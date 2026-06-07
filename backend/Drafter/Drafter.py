import logging
import os
from typing import Any, Dict, Optional, Set, Tuple

import uvicorn
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt
from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

from legal_draft import generate_legal_draft

load_dotenv()

logger = logging.getLogger(__name__)

app = FastAPI(title="DraftMate Drafter Service", version="2.0.0")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def create_content_control_run(paragraph, placeholder_tag: str, default_text: str):
    sdt = OxmlElement("w:sdt")

    sdt_pr = OxmlElement("w:sdtPr")

    tag = OxmlElement("w:tag")
    tag.set(qn("w:val"), placeholder_tag)
    sdt_pr.append(tag)

    alias = OxmlElement("w:alias")
    alias.set(qn("w:val"), placeholder_tag)
    sdt_pr.append(alias)

    showing = OxmlElement("w:showingPlcHdr")
    sdt_pr.append(showing)

    sdt.append(sdt_pr)

    sdt_content = OxmlElement("w:sdtContent")

    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:cs"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "Times New Roman")
    r_pr.append(r_fonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")
    r_pr.append(sz)

    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), "24")
    r_pr.append(sz_cs)

    r.append(r_pr)

    t = OxmlElement("w:t")
    t.text = f"[{default_text}]"
    r.append(t)

    sdt_content.append(r)
    sdt.append(sdt_content)

    paragraph._p.append(sdt)


def _token_core_segments(token: str) -> Tuple[str, str, str]:
    if not token:
        return "", "", ""
    start = 0
    while start < len(token) and not (token[start].isalnum() or token[start] == "_"):
        start += 1
    end = len(token)
    while end > start and not (token[end - 1].isalnum() or token[end - 1] == "_"):
        end -= 1
    return token[:start], token[start:end], token[end:]


def _apply_run_style(run, font_size: Pt, bold: bool):
    run.font.name = "Times New Roman"
    run.font.size = font_size
    run.bold = bold


def build_docx_with_controls(ai_data: dict, file_target_name: str) -> str:
    shared_storage_path = os.getenv("SHARED_STORAGE_PATH")
    if not shared_storage_path:
        raise ValueError("SHARED_STORAGE_PATH is not set.")

    os.makedirs(shared_storage_path, exist_ok=True)

    safe_name = (file_target_name or "").strip() or "draftmate_draft.docx"
    if not safe_name.lower().endswith(".docx"):
        safe_name = safe_name + ".docx"
    output_path = os.path.join(shared_storage_path, safe_name)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    metadata = ai_data.get("metadata") or {}
    placeholders_list = metadata.get("placeholders_detected") or []
    placeholders: Set[str] = set(x for x in placeholders_list if isinstance(x, str))

    content_blocks = ai_data.get("content") or []
    for block in content_blocks:
        if not isinstance(block, dict):
            continue

        element_type = block.get("element_type")
        text = block.get("text") or ""
        if not isinstance(text, str):
            text = str(text)

        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.line_spacing = 1.5

        if element_type == "header_block":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            font_size = Pt(14)
            bold = True
        elif element_type == "paragraph":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            font_size = Pt(12)
            bold = False
        elif element_type == "heading_1":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            font_size = Pt(13)
            bold = True
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            font_size = Pt(12)
            bold = False

        tokens = text.split(" ")
        for idx, token in enumerate(tokens):
            if idx > 0:
                space_run = paragraph.add_run(" ")
                _apply_run_style(space_run, font_size=font_size, bold=bold)

            prefix, core, suffix = _token_core_segments(token)
            if core and core in placeholders:
                if prefix:
                    r = paragraph.add_run(prefix)
                    _apply_run_style(r, font_size=font_size, bold=bold)

                create_content_control_run(paragraph, placeholder_tag=core, default_text=core)

                if suffix:
                    r = paragraph.add_run(suffix)
                    _apply_run_style(r, font_size=font_size, bold=bold)
            else:
                r = paragraph.add_run(token)
                _apply_run_style(r, font_size=font_size, bold=bold)

    doc.save(output_path)
    return output_path


class DraftCompileRequest(BaseModel):
    case_context: str
    legal_documents: Optional[str] = None
    document_type: str = "Legal Document"
    file_target_name: str = "draftmate_draft.docx"


@app.get("/")
def root():
    return {"service": "drafter-service", "status": "ok"}


@app.post("/v2/draft/compile")
def compile_draft(request: DraftCompileRequest):
    try:
        ai_data = generate_legal_draft(
            case_context=request.case_context,
            legal_documents=request.legal_documents,
            document_type=request.document_type,
        )
        output_path = build_docx_with_controls(ai_data=ai_data, file_target_name=request.file_target_name)
        return {
            "title": ai_data.get("title"),
            "metadata": ai_data.get("metadata"),
            "output_path": output_path,
            "file_name": os.path.basename(output_path),
        }
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.exception("Draft compilation failed.")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/v2/draft/forcesave")
def onlyoffice_forcesave(event: Dict[str, Any]):
    logger.info("OnlyOffice forcesave event received.")
    return {"error": 0}


@app.post("/v2/draft/callback")
def onlyoffice_callback(event: Dict[str, Any]):
    logger.info("OnlyOffice callback received.")
    return {"error": 0}


if __name__ == "__main__":
    uvicorn.run("Drafter:app", host="0.0.0.0", port=8003, reload=True)
